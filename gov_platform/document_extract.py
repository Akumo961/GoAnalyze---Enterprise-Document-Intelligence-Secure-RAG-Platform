"""Safe local document extraction with explicit OCR integration.

Text extraction is performed from supplied bytes. OCR is an opt-in executable
integration and is never treated as authoritative. Unsupported or malformed
files fail closed.
"""
from __future__ import annotations

import dataclasses
import hashlib
import io
import os
import subprocess
import tempfile


MAX_DOCUMENT_BYTES = 50 * 1024 * 1024
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "text/plain",
    "text/csv",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


@dataclasses.dataclass(frozen=True)
class ExtractedDocument:
    sha256: str
    text: str
    method: str
    page_count: int | None


def validate_upload(data: bytes, content_type: str) -> str:
    if not data:
        raise ValueError("empty_document")
    if len(data) > MAX_DOCUMENT_BYTES:
        raise ValueError("document_too_large")
    if content_type not in ALLOWED_CONTENT_TYPES:
        raise ValueError("unsupported_content_type")
    return hashlib.sha256(data).hexdigest()


def extract_text(data: bytes, content_type: str) -> ExtractedDocument:
    digest = validate_upload(data, content_type)
    if content_type.startswith("text/"):
        return ExtractedDocument(digest, data.decode("utf-8", errors="strict"), "text", None)
    if content_type == "application/pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data), strict=False)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return ExtractedDocument(digest, text, "pdf-text", len(reader.pages))
        except Exception as exc:
            raise ValueError("pdf_extraction_failed") from exc
    if content_type.endswith("wordprocessingml.document"):
        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs)
            return ExtractedDocument(digest, text, "docx-text", None)
        except Exception as exc:
            raise ValueError("docx_extraction_failed") from exc
    raise ValueError("unsupported_content_type")


def tesseract_ocr(data: bytes, *, timeout_seconds: int = 30) -> str:
    """Optional OCR integration. The caller must provide a trusted container.

    The input is written to a temporary file with a random OS-generated name;
    no user-controlled filename reaches the subprocess. Tesseract must already
    be installed and isolated by the deployment environment.
    """
    validate_upload(data, "application/pdf")
    with tempfile.TemporaryDirectory(prefix="goanalyze-ocr-") as directory:
        source = os.path.join(directory, "input.pdf")
        with open(source, "wb") as handle:
            handle.write(data)
        try:
            completed = subprocess.run(
                ["tesseract", source, "stdout", "--psm", "3"],
                capture_output=True,
                text=True,
                timeout=timeout_seconds,
                check=False,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise RuntimeError("ocr_unavailable_or_timeout") from exc
        if completed.returncode != 0:
            raise RuntimeError("ocr_failed")
        return completed.stdout
